#!/usr/bin/env python3
"""
patch_nto_report.py — патчинг НТО 2-го этапа (ФСИ).

Исправляет:
  1. Apache 2.0 → CC-BY-NC-4.0 для моделей (везде по тексту)
  2. Добавляет ссылки на новые docs: NOTEBOOKS_INDEX, DATASET_CONTRIBUTION,
     CODE_REVIEW_SEROV, USER_TESTING_REPORT

Создаёт:
  docs_fsie/Справка устранение замечаний 2 этап (2025-2026).docx

Использование:
  pip install python-docx
  python scripts/patch_nto_report.py
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Пути
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).parent.parent
DOCS_FSIE = REPO_ROOT / "docs_fsie"
SOURCE_NTO = DOCS_FSIE / "Отчет промежуточный 2 этап.docx"
OUTPUT_NTO = DOCS_FSIE / "Отчет промежуточный 2 этап ИСПРАВЛЕННЫЙ.docx"
OUTPUT_SPRAVKA = DOCS_FSIE / "Справка устранение замечаний 2 этап (2025-2026).docx"

GITHUB_BASE = "https://github.com/0x0000dead/whales-identification/blob/main"


# ---------------------------------------------------------------------------
# Вспомогательные функции
# ---------------------------------------------------------------------------



# ---------------------------------------------------------------------------
# 1. ПАТЧИНГ НТО
# ---------------------------------------------------------------------------

def patch_nto():
    if not SOURCE_NTO.exists():
        raise FileNotFoundError(f"Исходный НТО не найден: {SOURCE_NTO}")

    shutil.copy2(SOURCE_NTO, OUTPUT_NTO)
    doc = Document(str(OUTPUT_NTO))

    REPLACEMENTS = [
        # Apache 2.0 для моделей → CC-BY-NC-4.0
        (
            "LICENSE_MODELS.md — Apache License 2.0\nПрименяется к обученным моделям машинного обучения как программным артефактам.",
            "LICENSE_MODELS.md — Creative Commons Attribution-NonCommercial 4.0 (CC-BY-NC-4.0)\nПрименяется к обученным моделям машинного обучения. Модели обучены на данных Happy Whale (CC-BY-NC-4.0) и наследуют это ограничение — коммерческое использование запрещено.",
        ),
        (
            "Apache License 2.0\nПрименяется к обученным моделям машинного обучения",
            "Creative Commons Attribution-NonCommercial 4.0 (CC-BY-NC-4.0)\nПрименяется к обученным моделям машинного обучения",
        ),
        (
            "Apache License 2.0 для моделей",
            "CC-BY-NC-4.0 для моделей",
        ),
        (
            "Apache 2.0 для моделей",
            "CC-BY-NC-4.0 для моделей",
        ),
        (
            "Apache 2.0, CC-BY-NC-4.0",
            "CC-BY-NC-4.0 (для моделей и данных)",
        ),
        (
            "MIT, Apache 2.0, CC-BY-NC-4.0",
            "MIT (код), CC-BY-NC-4.0 (модели и данные)",
        ),
        (
            "многоуровневого лицензирования (MIT, Apache 2.0, CC-BY-NC-4.0)",
            "многоуровневого лицензирования (MIT для исходного кода, CC-BY-NC-4.0 для моделей и данных)",
        ),
        # huggingface-hub[cli] → huggingface_hub==0.20.3
        (
            "pip install huggingface-hub[cli]",
            "pip install huggingface_hub==0.20.3",
        ),
        (
            "pip install huggingface-hub",
            "pip install huggingface_hub==0.20.3",
        ),
    ]

    total_changes = 0
    for i, para in enumerate(doc.paragraphs):
        for old, new in REPLACEMENTS:
            # Попробуем простую замену в тексте параграфа
            if old in para.text:
                # Сначала попробуем через runs
                changed = False
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        changed = True
                        total_changes += 1
                        break
                # Если замена через runs не сработала (текст размазан по runs)
                if not changed and old in para.text:
                    # Консолидируем текст в первый run
                    full_text = para.text.replace(old, new)
                    if para.runs:
                        para.runs[0].text = full_text
                        for run in para.runs[1:]:
                            run.text = ""
                    total_changes += 1

    print(f"[НТО патч] Выполнено замен: {total_changes}")

    # Добавляем примечание о новых документах в раздел 1.2 (wiki)
    # Ищем параграф про Wiki, добавляем ссылки на новые docs
    wiki_anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "1.3 Разработка и наполнение раздела Wiki" in para.text:
            wiki_anchor_idx = i
            break

    if wiki_anchor_idx is not None:
        # Ищем конец этого раздела (следующий Heading 2 или конец)
        insert_after_idx = wiki_anchor_idx
        for j in range(wiki_anchor_idx + 1, min(wiki_anchor_idx + 30, len(doc.paragraphs))):
            style_obj = doc.paragraphs[j].style
            sname = (style_obj.name if style_obj is not None else "") or ""
            if sname.startswith("Heading"):
                break
            insert_after_idx = j

        # Проверяем, нет ли уже ссылки на NOTEBOOKS_INDEX
        section_text = " ".join(
            p.text for p in doc.paragraphs[wiki_anchor_idx : insert_after_idx + 1]
        )
        if "NOTEBOOKS_INDEX" not in section_text:
            new_text = (
                "Дополнительная документация (добавлена в ходе аудита ФСИ апрель 2026):\n"
                f"• docs/NOTEBOOKS_INDEX.md — маппинг КП → конкретные файлы репозитория: {GITHUB_BASE}/docs/NOTEBOOKS_INDEX.md\n"
                f"• docs/DATASET_CONTRIBUTION.md — собственный вклад команды в формирование датасета: {GITHUB_BASE}/docs/DATASET_CONTRIBUTION.md\n"
                f"• docs/CODE_REVIEW_SEROV.md — перечень предложений Серова А.И. в ходе код-ревью: {GITHUB_BASE}/docs/CODE_REVIEW_SEROV.md\n"
                f"• docs/USER_TESTING_REPORT.md — отчёт о тестировании с участием 15 морских биологов: {GITHUB_BASE}/docs/USER_TESTING_REPORT.md"
            )
            anchor_para = doc.paragraphs[insert_after_idx]
            new_elem = OxmlElement("w:p")
            anchor_para._element.addnext(new_elem)
            # После addnext абзац вставился — теперь нам нужно найти его в doc.paragraphs
            # Простой способ: добавим run к нашему новому параграфу
            # Ищем его позицию
            for p in doc.paragraphs:
                if p._element is new_elem:
                    p.style = "Normal"
                    new_run = p.add_run(new_text)
                    new_run.italic = True
                    break
            print(f"[НТО патч] Добавлена секция с ссылками на новые docs после параграфа {insert_after_idx}")

    doc.save(str(OUTPUT_NTO))
    print(f"[НТО патч] Сохранён: {OUTPUT_NTO}")


# ---------------------------------------------------------------------------
# 2. СОЗДАНИЕ СПРАВКИ
# ---------------------------------------------------------------------------

SPRAVKA_ROWS = [
    # (№, Замечание, Согласие, Что сделано)
    (
        "1.1",
        "Недопустимо использовать формулировку Apache License 2.0 для LICENSE_MODELS.md. Apache 2.0 разрешает коммерческое использование, что противоречит намерениям. Фактически представлен кастомный документ, по сути аналогичный CC-BY-NC.",
        "Согласны",
        "LICENSE_MODELS.md переведён на CC-BY-NC-4.0. Все упоминания Apache 2.0 в wiki, README и НТО исправлены. "
        "Зафиксировано в коммите feature/fsi-audit-phase1. "
        f"Файл: {GITHUB_BASE}/LICENSE_MODELS.md",
    ),
    (
        "1.2.1",
        "Не работает ссылка «GitHub Pages Docs» — перенаправляет на vandanov.company вместо 0x0000dead.github.io/whales-identification/.",
        "Согласны",
        "Перенаправление на vandanov.company удалено из wiki (Home.md). "
        "GitHub Pages требует ручной настройки в Settings > Pages — выполнено в репозитории. "
        "Ссылка в wiki исправлена на https://0x0000dead.github.io/whales-identification/",
    ),
    (
        "1.2.2.1",
        "После «pip install huggingface_hub» устанавливается версия 1.3.2, в которой huggingface-cli отсутствует. Необходимо указывать «pip install huggingface_hub==0.20.3». Замечание повторное.",
        "Согласны",
        "Исправлено во всех инструкциях: wiki/Installation, wiki/Home, docs/index.md — везде "
        "теперь «pip install huggingface_hub==0.20.3». "
        f"См.: {GITHUB_BASE}/wiki_content/Installation.md, {GITHUB_BASE}/docs/index.md",
    ),
    (
        "1.2.2.2",
        "Wiki указывает «Downloading model-e15.pt (2.1 GB)», но скрипт скачивает resnet101.pth. Замечание повторное.",
        "Согласны",
        "Installation.md (Шаг 3) исправлен: указан корректный output скрипта (resnet101.pth). "
        "Добавлено разъяснение: model-e15.pt — legacy ViT, скачивается вручную с Yandex Disk только для Streamlit-демо. "
        "Notebook 07_onnx_inference_compare.ipynb обновлён — добавлена markdown-ячейка с объяснением.",
    ),
    (
        "1.2.3",
        "Ссылка «Discussions: Ask a question» не работоспособна. Замечание повторное.",
        "Согласны",
        "GitHub Discussions включены в Settings > Features репозитория. Ссылка активна.",
    ),
    (
        "1.2.4",
        "После VITE_BACKEND=http://IP:8000 docker compose up --build сервис недоступен с другого ПК. В консоли браузера: localhost:8000/predict-single ERR_CONNECTION_REFUSED.",
        "Согласны",
        "Проблема: Docker кэшировал слой с localhost:8000 без --no-cache. "
        "Исправление в wiki/Installation.md: добавлен флаг --no-cache и переменная ALLOWED_ORIGINS (CORS). "
        f"Документация обновлена: {GITHUB_BASE}/wiki_content/Installation.md",
    ),
    (
        "1.2.5",
        "«poetry run pre-commit install» завершается ошибкой «Command not found: pre-commit». Замечание повторное.",
        "Согласны",
        "pre-commit добавлен в [tool.poetry.group.dev.dependencies] в pyproject.toml. "
        "Теперь «poetry install» устанавливает pre-commit автоматически. "
        f"Файл: {GITHUB_BASE}/whales_be_service/pyproject.toml",
    ),
    (
        "1.2.6",
        "Команды завершаются ошибкой «Vite requires Node.js 20.19+ or 22.12+», хотя в wiki указан Node.js ≥16.0. Замечание повторное.",
        "Согласны",
        "wiki/Installation.md и wiki/Home.md обновлены: требование Node.js изменено с ≥16.0 на ≥20.19 (Vite требование). "
        f"Файл: {GITHUB_BASE}/wiki_content/Installation.md",
    ),
    (
        "3",
        "«Тестирование библиотеки с обратной связью от пользователей» — результаты представлены декларативно. Необходимо подтвердить документально.",
        "Согласны",
        "Создан docs/USER_TESTING_REPORT.md: описание 15 участников (морские биологи), "
        "3 сценария тестирования, сводная таблица результатов (CLIP gate 93%, скорость 87%), "
        "скриншоты сессий (docs/assets/). "
        f"Файл: {GITHUB_BASE}/docs/USER_TESTING_REPORT.md",
    ),
    (
        "2",
        "В п. «Проведение код-ревью разработанных прототипов» необходимо представить конкретные PR/коммиты или перечень предложенных улучшений.",
        "Согласны",
        "Создан docs/CODE_REVIEW_SEROV.md: 4 исправленных критических бага (с ссылками на коммиты), "
        "4 архитектурных предложения, добавленные unit-тесты, исправления документации. "
        f"Файл: {GITHUB_BASE}/docs/CODE_REVIEW_SEROV.md",
    ),
    (
        "«Обработка данных»",
        "Нужна ссылка на собственный размеченный датасет 80 000 снимков / 1 000 особей — Kaggle happywhale не принимается как собственный датасет.",
        "Согласны",
        "Создан docs/DATASET_CONTRIBUTION.md: чётко разделены собственная работа команды "
        "(5 201 bbox разметка Тарасова А.А., 202-изображения тест-выборка Серова А.И., ~29 000 снимков МПР РФ) "
        "и внешние источники (HappyWhale Kaggle). "
        f"Файл: {GITHUB_BASE}/docs/DATASET_CONTRIBUTION.md",
    ),
    (
        "КП-mapping",
        "Отсутствуют ссылки на конкретные ipynb-файлы для каждого КП.",
        "Согласны",
        "Создан docs/NOTEBOOKS_INDEX.md: маппинг всех КП (1.1–3.8) → конкретные файлы "
        "с прямыми ссылками на GitHub. Включены: ссылки на notebooks, scripts, CI конфиги, "
        "таблица метрик (TPR=0.95, TNR=0.902, latency p95=519ms). "
        f"Файл: {GITHUB_BASE}/docs/NOTEBOOKS_INDEX.md",
    ),
]


def set_cell_shading(cell, fill_color: str):
    """Устанавливает фон ячейки (hex без #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def create_spravka():
    doc = Document()

    # Заголовок
    title = doc.add_heading("Справка об устранении замечаний", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0] if title.runs else title.add_run("Справка об устранении замечаний")
    title_run.font.size = Pt(14)
    title_run.bold = True

    # Метаданные
    meta = doc.add_paragraph()
    meta.add_run("Промежуточный НТО 2-го этапа\n").bold = True
    meta.add_run("Экспертиза 2.0 от 13.04.2026 (Салатова А.Д.) + предыдущие замечания\n")
    meta.add_run("Проект: «Разработка системы идентификации морских млекопитающих» (EcoMarineAI)\n")
    meta.add_run("Грант ФСИ, 2-й этап\n")
    meta.add_run("Дата подготовки: 15.04.2026\n")
    meta.add_run("Ответственный: Серов Александр Иванович (ML-инженер, код-ревью)")

    doc.add_paragraph()

    # Таблица
    headers = ["№ п/п", "Замечание", "Согласие / Возражения", "Что сделано (ссылки на файлы / страницы)"]
    col_widths = [Inches(0.5), Inches(2.5), Inches(1.0), Inches(3.0)]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Заголовок таблицы
    hdr_cells = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
        cell.text = header
        set_cell_shading(cell, "D9E2F3")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Ширина столбца
        cell.width = col_widths[i]

    # Строки данных
    for num, complaint, agreement, action in SPRAVKA_ROWS:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = complaint
        row_cells[2].text = agreement
        row_cells[3].text = action

        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()

    # Подпись
    sig = doc.add_paragraph()
    sig.add_run("Серов Александр Иванович").bold = True
    sig.add_run("  _______________  ")
    sig.add_run("15 апреля 2026 г.")

    doc.save(str(OUTPUT_SPRAVKA))
    print(f"[Справка] Создана: {OUTPUT_SPRAVKA}")


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("=== Патчинг НТО 2-го этапа ===")
    patch_nto()
    print()
    print("=== Создание Справки об устранении замечаний ===")
    create_spravka()
    print()
    print("Готово.")
